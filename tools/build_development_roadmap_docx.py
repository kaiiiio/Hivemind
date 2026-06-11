from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "HIVEMIND_Development_Roadmap.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D9E2EC", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.12


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_header_row(table, labels, fill="F4F6F9"):
    row = table.rows[0]
    set_repeat_table_header(row)
    for idx, label in enumerate(labels):
        cell = row.cells[idx]
        set_cell_shading(cell, fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(31, 77, 120)
        run.font.size = Pt(9.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    add_header_row(table, headers)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor(31, 41, 55)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(item)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, "F4F6F9")
    set_cell_border(cell, color="C9D6E2")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    p2.add_run(body)
    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(31, 41, 55)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("HIVEMIND Development Roadmap")
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor(107, 114, 128)


def build_doc():
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("HIVEMIND Development Roadmap")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.add_run(
        "Functional MVP roadmap for an AI-swarm capital-markets intelligence terminal with 10-15 agents."
    ).italic = True

    meta_rows = [
        ("Prepared for", "Founder planning, grant applications, investor discussions, and developer execution"),
        ("Roadmap horizon", "0-180 days, with a grant-funded 90-day MVP path"),
        ("Product scope", "Indian capital markets: equities, debt/rates, FX, commodities, derivatives context, policy, tenders, filings, price action"),
        ("MVP constraint", "End-of-day and historical intelligence first; no automated trading or exchange-grade live feed in MVP"),
    ]
    add_table(doc, ["Field", "Roadmap Definition"], meta_rows, [Inches(1.65), Inches(4.85)])

    add_callout(
        doc,
        "Roadmap principle",
        "Do not build a generic stock screener. Build an evidence-first terminal where every situation brief can be traced back to raw sources, parsed records, features, agent outputs, and later outcome labels.",
    )

    doc.add_heading("1. Product Target", level=1)
    doc.add_paragraph(
        "The end product is a professional capital-markets intelligence terminal. The user should be able to search across issuers, securities, sectors, commodities, rates, FX, derivatives context, policy, tenders, filings, price action, and portfolio exposure. The system should convert those fragments into evidence-cited market situations."
    )
    add_bullets(
        doc,
        [
            "Primary workflow: detect, investigate, validate, brief, monitor, and replay market situations.",
            "Primary user value: find important changes earlier, understand why they matter, and avoid unsupported narratives.",
            "Primary technical moat: evidence archive, market graph, agent workflows, outcome memory, and source-health know-how.",
        ],
    )

    doc.add_heading("2. MVP Boundaries", level=1)
    add_table(
        doc,
        ["In Scope For MVP", "Out Of Scope For MVP"],
        [
            ("BSE/NSE EOD and historical data ingestion", "Tick-by-tick or exchange-grade real-time feed"),
            ("Official filings, announcements, policy/tender sources", "Automated execution or portfolio order placement"),
            ("10-15 model-agnostic agents", "Running 30-40 agents on every event"),
            ("Evidence IDs, parsed records, graph edges, and situation briefs", "Unverifiable AI-generated facts"),
            ("Event-study replay and false-positive labels", "Guaranteed alpha or performance claims"),
            ("Prototype terminal UI and investor demo", "Full commercial Bloomberg-like coverage"),
        ],
        [Inches(3.25), Inches(3.25)],
    )

    doc.add_heading("3. Phase Roadmap", level=1)
    add_table(
        doc,
        ["Phase", "Duration", "Objective", "Exit Criteria"],
        [
            ("0. Architecture and grant readiness", "Week 0-1", "Finalize canonical docs, investor whitepaper site, budget, and source map.", "Grant-ready roadmap, cost estimate, deployable site, initial backlog."),
            ("1. Evidence spine MVP", "Weeks 1-4", "Build raw evidence lake, EOD data ingestion, filing ingestion, and instrument identity.", "Raw evidence stored with IDs; first historical price and filing backfill works."),
            ("2. Situation intelligence alpha", "Weeks 5-8", "Add 10-15 routed agents and first situation engines.", "Order wins, results, policy/tender, and unexplained price-action situations produce evidence-cited briefs."),
            ("3. Terminal private beta", "Weeks 9-12", "Ship terminal UI for search, situation queue, evidence explorer, agent traces, and watchlists.", "User can investigate a situation end to end from terminal screen."),
            ("4. Validation and portfolio layer", "Days 90-150", "Add event-study replay, outcome memory, portfolio exposure, and post-mortem loops.", "Alerts can be scored as early, late, false positive, useful watch, or missed catalyst."),
            ("5. Multi-asset expansion", "Days 150-180+", "Expand debt/rates, FX, commodities, derivatives context, and specialist agent pools.", "Cross-asset situation engines and paid-data pilots are justified by measured usefulness."),
        ],
        [Inches(1.35), Inches(0.95), Inches(2.05), Inches(2.15)],
    )

    doc.add_heading("4. 90-Day MVP Sprint Plan", level=1)
    add_table(
        doc,
        ["Sprint", "Build Focus", "Core Deliverables"],
        [
            ("Sprint 1", "Repository, schema, and data identity", "Project structure, Postgres schema, instrument universe, BSE/NSE mappings, raw evidence table, source registry."),
            ("Sprint 2", "Official source ingestion", "BSE/NSE filings ingestion, EOD OHLCV/delivery backfill, checksum storage, scheduler, retry logic."),
            ("Sprint 3", "Parsing and entity resolution", "Filing parser, tender/policy parser, entity resolver, alias maps, parser confidence and evidence spans."),
            ("Sprint 4", "Feature store and graph", "Price/volume/liquidity features, event windows, company-sector-policy-tender graph, graph writer worker."),
            ("Sprint 5", "Agent runtime foundation", "Provider interface, model router, agent registry, tool gateway, memory broker, structured JSON output contracts."),
            ("Sprint 6", "First situation engines", "Order win, results, promoter/capex, policy/tender tailwind, unexplained price-action candidate generation."),
            ("Sprint 7", "Research job orchestration", "Prompt-to-job compiler, source discovery agent, hard filters, debate loop, synthesis, rejection reasons."),
            ("Sprint 8", "Terminal MVP", "Search, situation monitor, situation detail, evidence explorer, agent trace panel, watchlist workflow."),
            ("Sprint 9", "Replay and demo hardening", "Event-study runner, outcome labels, smoke tests, seed examples, demo walkthrough, deployment checklist."),
        ],
        [Inches(0.9), Inches(1.75), Inches(3.85)],
    )

    doc.add_heading("5. Technical Workstreams", level=1)
    add_table(
        doc,
        ["Workstream", "Responsibilities", "MVP Acceptance Checks"],
        [
            ("Data foundation", "Source adapters, raw evidence storage, parser outputs, canonical IDs, backfill jobs.", "Every derived object can trace back to raw evidence IDs and timestamps."),
            ("AI swarm runtime", "Provider abstraction, agent manifests, routing, memory packs, structured outputs, critique loops.", "Agents can run on interchangeable providers and produce schema-valid outputs."),
            ("Market graph and memory", "Graph edges, vector retrieval, outcome memory, agent run logs, source reliability.", "Situation pages retrieve relevant evidence, graph neighbors, and prior outcomes."),
            ("Quant validation", "Event-study windows, abnormal returns, liquidity/slippage assumptions, false-positive labels.", "At least 20-30 historical cases can be replayed and scored."),
            ("Terminal frontend", "Search, situation queue, evidence explorer, agent traces, watchlists, portfolio notes.", "A user can inspect a situation without reading raw logs or code."),
            ("DevOps and QA", "Docker/dev setup, scheduled jobs, observability, backups, deployment, smoke tests.", "Daily ingestion health, failed source alerts, and reproducible demo environment exist."),
        ],
        [Inches(1.45), Inches(2.55), Inches(2.5)],
    )

    doc.add_heading("6. Agent Rollout Plan", level=1)
    add_table(
        doc,
        ["Wave", "Agents", "Purpose"],
        [
            ("Wave 1", "Source Health, Evidence Intake, Filing Parser, Entity Resolver", "Make ingestion reliable before reasoning starts."),
            ("Wave 2", "Market Data Worker, Price/Volume Agent, Graph Builder, Macro/Policy Agent", "Create market context and relationship memory."),
            ("Wave 3", "Search Investigator, Tender Agent, Valuation Agent, Quant Validator", "Fill missing context, test materiality, and validate market behavior."),
            ("Wave 4", "Bull Case, Risk Review, Synthesis", "Create balanced evidence-cited situation briefs with uncertainty and rejection logic."),
        ],
        [Inches(1.0), Inches(2.6), Inches(2.9)],
    )
    add_callout(
        doc,
        "Routing rule",
        "Do not run every agent on every event. The router should select agents based on situation family, data freshness, missing fields, portfolio exposure, source confidence, and model budget.",
    )

    doc.add_heading("7. Terminal Feature Roadmap", level=1)
    add_table(
        doc,
        ["Feature", "MVP Version", "Later Version"],
        [
            ("Universal search", "Issuer, security, event, source, and situation search.", "Cross-asset semantic search with graph neighborhoods and saved queries."),
            ("Situation monitor", "Ranked queue with evidence status, freshness, materiality, and uncertainty.", "Personalized ranking by portfolio exposure and user preferences."),
            ("Situation detail", "Trigger, evidence, materiality, price action, valuation, risks, agent outputs.", "Scenario analysis, hedge context, and similar-case replay."),
            ("Evidence explorer", "Raw evidence IDs, source links, parser fields, timestamps.", "Source reliability analytics and parser comparison history."),
            ("Portfolio workspace", "Manual watchlist and notes.", "Broker import, exposure map, risk budgets, post-mortems, and hedge suggestions."),
            ("Research jobs", "Prompt-to-job workflow for structured investigations.", "Reusable templates, scheduled research, multi-agent debate scorecards."),
        ],
        [Inches(1.65), Inches(2.45), Inches(2.4)],
    )

    doc.add_heading("8. Data And Infrastructure Roadmap", level=1)
    add_numbered(
        doc,
        [
            "Start with Postgres as the system of record: instruments, evidence, parsed records, events, features, agent runs, and outcomes.",
            "Use object storage for raw PDFs, HTML snapshots, CSVs, API payloads, and search snapshots.",
            "Use pgvector first for semantic retrieval; add a dedicated vector store only if scale or performance demands it.",
            "Represent the graph initially in relational edge tables; move to Neo4j/Memgraph only after graph queries become central.",
            "Run ingestion and agent jobs through a queue so source polling, parsing, research, and replay can be retried independently.",
            "Keep Vercel for static/investor site and terminal frontend previews; deploy API/workers separately on VPS or cloud containers.",
        ],
    )

    doc.add_heading("9. Quality Gates", level=1)
    add_table(
        doc,
        ["Gate", "Requirement"],
        [
            ("Evidence gate", "No situation brief is published without source IDs, timestamps, and evidence lineage."),
            ("Parser gate", "Parser outputs must include confidence and raw evidence spans for important fields."),
            ("Agent gate", "Agent outputs must be valid JSON with claims, evidence IDs, uncertainty, and missing facts."),
            ("Risk gate", "Risk Review Agent must run before synthesis for investable situations."),
            ("Replay gate", "Every published situation should later receive an outcome label."),
            ("Cost gate", "LLM calls must log provider, model, tokens, cost estimate, and context pack ID."),
            ("Deployment gate", "Main branch deploys cleanly; staging/demo branch remains reproducible."),
        ],
        [Inches(1.45), Inches(5.05)],
    )

    doc.add_heading("10. Team And Execution Model", level=1)
    doc.add_paragraph(
        "For a grant-funded MVP, founder engineering effort is assumed as in-kind contribution. Cash should be preserved for data, infrastructure, AI credits, testing, and deployment."
    )
    add_table(
        doc,
        ["Role", "MVP Responsibility"],
        [
            ("Founder / Product Lead", "Use-case selection, source priority, evaluation cases, investor/grant communication."),
            ("Backend/Data Engineer", "Adapters, schema, ingestion jobs, storage, queues, API services."),
            ("AI Engineer", "Agent manifests, model routing, structured outputs, memory broker, evaluation prompts."),
            ("Frontend Engineer", "Terminal UX, situation pages, evidence explorer, watchlist/portfolio workflows."),
            ("Quant/Research Reviewer", "Event-study assumptions, validation labels, false-positive review, financial sanity checks."),
        ],
        [Inches(1.8), Inches(4.7)],
    )

    doc.add_heading("11. Risks And Mitigation", level=1)
    add_table(
        doc,
        ["Risk", "Mitigation"],
        [
            ("Free sources are incomplete or change structure", "Store raw evidence, monitor source health, maintain adapters, reserve paid-data pilot budget."),
            ("AI hallucination contaminates facts", "AI cannot write facts without evidence IDs; unverifiable claims become EVIDENCE_GAP."),
            ("Too many agents increase cost and noise", "Use router, confidence gates, cheap extraction models, and post-run cost logging."),
            ("Historical replay is biased", "Use point-in-time data, corporate-action handling, explicit windows, and false-positive labels."),
            ("Terminal becomes generic dashboard", "Build around situation workflows: trigger, materiality, evidence, price action, valuation, risk, outcome."),
            ("MVP scope expands too far", "Keep exchange-grade live feed, automated execution, and full paid data coverage out of MVP."),
        ],
        [Inches(2.15), Inches(4.35)],
    )

    doc.add_heading("12. Definition Of Done For Functional MVP", level=1)
    add_bullets(
        doc,
        [
            "At least 3-4 situation families produce evidence-cited briefs from real historical and current data.",
            "At least 10 routed agents/workers run through the same model-agnostic interface.",
            "Every situation has raw evidence IDs, parsed records, generated features, and agent traces.",
            "Terminal UI supports search, situation queue, situation detail, evidence explorer, and watchlist notes.",
            "Event-study replay works on selected historical cases and writes outcome labels.",
            "Deployment is reproducible, with documented local setup and Vercel/frontend deployment path.",
        ],
    )

    doc.add_heading("13. Immediate Next Actions", level=1)
    add_numbered(
        doc,
        [
            "Commit and push the current whitepaper site and docs to GitHub.",
            "Confirm Vercel is connected to the GitHub repository and the deployed branch.",
            "Create the backend repository structure and database schema migrations.",
            "Implement source registry, raw evidence table, instrument identity, and first BSE/NSE ingestion jobs.",
            "Create agent manifest format and mock provider so agents can be tested before spending API credits.",
            "Seed 20-30 historical situations for validation and demo cases.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
